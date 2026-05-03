from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

from .base_parser import BaseParser, ParsedDocument

SUPPORTED_EXTENSIONS = {".docx"}


class DOCXParser(BaseParser):
    """
    Parses Word documents using python-docx.

    Extracts:
    - Paragraphs (body text)
    - Headings (with level info preserved in metadata)
    - Tables (converted to plain text rows)
    - Core document properties (author, title, etc.)

    Note: .doc (old binary format) is NOT supported — only .docx.
    If you need .doc support, add python-docx2txt or use LibreOffice headlessly.
    """

    def supports(self, file_path: str | Path) -> bool:
        return Path(file_path).suffix.lower() in SUPPORTED_EXTENSIONS

    def parse(self, file_path: str | Path) -> ParsedDocument:
        path = self._validate_file(file_path)

        if not self.supports(path):
            raise ValueError(f"DOCXParser does not support: {path.suffix}")

        doc = Document(str(path))
        sections = []      # logical sections (heading → paragraphs)
        all_text_parts = []

        current_section = {"heading": None, "level": 0, "text_parts": []}

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name or ""

            if style_name.startswith("Heading"):
                # Save the previous section before starting a new one
                if current_section["text_parts"]:
                    sections.append(self._flush_section(current_section))

                # Extract heading level (e.g. "Heading 1" → 1)
                try:
                    level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    level = 1

                current_section = {
                    "heading": text,
                    "level": level,
                    "text_parts": [text],
                }
            else:
                current_section["text_parts"].append(text)

            all_text_parts.append(text)

        # Flush the last section
        if current_section["text_parts"]:
            sections.append(self._flush_section(current_section))

        # Extract tables — each table becomes a text block
        for table in doc.tables:
            table_text = self._table_to_text(table)
            if table_text:
                all_text_parts.append(table_text)
                sections.append({
                    "page_num": len(sections) + 1,
                    "text": table_text,
                    "metadata": {"type": "table"},
                })

        full_text = "\n\n".join(all_text_parts)

        # Core properties (author, title, etc.)
        props = doc.core_properties
        metadata = {
            "author": props.author or "",
            "title": props.title or "",
            "created": str(props.created) if props.created else "",
            "modified": str(props.modified) if props.modified else "",
            "section_count": len(sections),
            "file_size_bytes": path.stat().st_size,
        }

        title = props.title or path.stem

        return ParsedDocument(
            source_path=str(path),
            file_type="docx",
            title=title,
            full_text=full_text,
            pages=sections,  # "pages" = logical sections for DOCX
            metadata=metadata,
        )

    def _flush_section(self, section: dict) -> dict:
        return {
            "page_num": None,  # DOCX has no page numbers without rendering
            "text": "\n".join(section["text_parts"]),
            "metadata": {
                "heading": section["heading"],
                "heading_level": section["level"],
                "type": "section",
            },
        }

    def _table_to_text(self, table) -> str:
        """Convert a docx Table object to a readable plain-text grid."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            # Remove duplicate adjacent cells (merged cells repeat in python-docx)
            deduped = [cells[0]] + [c for i, c in enumerate(cells[1:], 1) if c != cells[i-1]]
            rows.append(" | ".join(deduped))
        return "\n".join(rows)